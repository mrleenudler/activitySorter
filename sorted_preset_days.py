from os import dup
import pandas as pd
from collections import Counter
import random

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# def __init__

# NB! Skal all input gjøres til .lower() ?

# target Excelark bør være userInput

df = pd.read_excel('Påmelding til turer!.xlsx')
fulltNavn = df['Skriv inn fornavn og etternavn'].tolist() # unique Identifier
klasse = df['Hvilken klasse går du i?'].tolist() # Class
Aktivitet1 = df['Hvilke aktiviteter vil du være med på? Velg så mange du vil, så vil du bli satt opp på to aktiviteter.'].tolist() # Acting as wanted group
# aktivitet2 = df['Hvilken aktivitet vil du være med på tirsdag 11. juni. Velg en ny aktivitet, ikke den du valgte for mandag.'].tolist() # Acting as wanted group
#priority = df['Prioritet'] # Wishes are sorted by priority
#friends = [] # Used to group friends together

# Flere verdier kan være userInput -> GUI vurderes etter hvert
highPriorityString = "Dette har jeg veldig lyst til"
lowPriorityString = "Dette har jeg litt lyst til"
maxActivitiesPerStudent = 2 # Bør være userInput
# Bør hentes fra Excelark
maxApplicationsPerStudent = 14 # aka totalNumberOfActivities
# All available unique students
allePaameldteElever = set(fulltNavn) # Kanskje heller set(df['Fullt navn'].tolist())
#VIKTIG! første oppføring er 'nan'
# Index positions of Excel source workbook. Should be user input or pulled automatically from workbook
inputTupleNamePosition = 9
inputTupleClassPosition = 14
inputTupleActivityPosition = 17
# Index positions for generated
applicationTupleNamePosition = 0
applicationTupleClassPosition = 1
applicationTupleActivityPosition = 2
applicationTupleDayPosition = 4
applicationTuplePriorityPosition = 3

# (navn, klasse, aktivitet, prioritet, dag)

#print("DEBUG: allePaameldteElever ", allePaameldteElever)
# Creating activities with max number of participants
# NB! Bør hentes fra excel/dataframe -> Komplisert, valgene er ikke tilgjengelige som en enhet fra Excel.
#fordeling = {"emel": [], "anne marie": [], "sveinung": [], "natasha": [], "elisabeth": [], "andreas": []}
fordeling = {"Orientering": [], "Frisbee Ultimate": [], "Kino i hall 2": [], \
    "Langtur i Østmarka": [], "Python nybegynner": [], "Brettspilldag": [], \
        "Ta med bok og les - det servers drikke!": [], "Python for de som kan litt fra før": [], \
            "Trigonometri": [], "Språkkafé": [], "Tur i Botanisk Hage + utendørsklatring i buldreveggen på Enerhaugen": [], \
                "Basket i Kubaparken": [], "Joggetur rundt Maridalsvannet": [], "Piknik i Torshovparken": []}
fordelingMax = {"Orientering": 1000, "Frisbee Ultimate": 1000, "Kino i hall 2": 1000, \
    "Langtur i Østmarka": 1000, "Python nybegynner": 1000, "Brettspilldag": 1000, \
        "Ta med bok og les - det servers drikke!": 1000, "Python for de som kan litt fra før": 1000, \
            "Trigonometri": 1000, "Språkkafé": 1000, "Tur i Botanisk Hage + utendørsklatring i buldreveggen på Enerhaugen": 1000, \
                "Basket i Kubaparken": 1000, "Joggetur rundt Maridalsvannet": 1000, "Piknik i Torshovparken": 1000}
fordelingDag = {"Orientering": "tirsdag", "Frisbee Ultimate": "mandag", "Kino i hall 2": "mandag", \
    "Langtur i Østmarka": "mandag", "Python nybegynner": "mandag", "Brettspilldag": "tirsdag", \
        "Ta med bok og les - det servers drikke!": "tirsdag", "Python for de som kan litt fra før": "mandag", \
            "Trigonometri": "mandag",  "Språkkafé": "tirsdag", "Tur i Botanisk Hage + utendørsklatring i buldreveggen på Enerhaugen": "tirsdag", \
                "Basket i Kubaparken": "mandag", "Joggetur rundt Maridalsvannet": "tirsdag", "Piknik i Torshovparken": "tirsdag"}
# eleverMedBekreftede aktiviteter må ta en liste over dager de er opptatt.



# NOTAT: Juster ned anbefalt linjeavstand i word-filen

# DEBUG - funker
# for key in fordeling:
#     fordeling[key] = []
#     print("Cleaning")
# print(fordeling["Orientering"])
# for item in fordeling:
#     print("Exists?", item, len(item))


# VIKTIG - bør kanskje fordele aktiviteter på dager? -> DONE!
# NB! Bør være userInput
#fordelingMax = {"emel": 120, "anne marie": 30, "sveinung": 25, "natasha": 50, "elisabeth": 8, "andreas": 100}
#fordelingMax = {"Orientering": 100, "Frisbee Ultimate": 100, "Kino i hall 2": 100, "Langtur i Østmarka": 100, "Python nybegynner": 100, "Brettspilldag": 100}
# Keeping track of number of activities each student is assigned to
# kanskje allePaamledteElever er redundant ? 

# Pass på at index til aktivitetene er riktig med ny layout på Forms - DONE
def structure_applications(inputApplications):
    # KAN LEGGE TIL KODE FOR Å MERGE DUPLICATES!
    # print(inputApplications[0])
    # print(inputApplications[0][9], inputApplications[0][14], inputApplications[0][17])
    allApplications = []
    for tup in inputApplications:
        # Iterating through chosen activities
        for aktivitet1 in tup[inputTupleActivityPosition].split(';'):
            if aktivitet1: # Fixing trailing ';' (perhaps)
                aktivitetsdag = ""
                # Hack to account for  "Kino i hall 2" running two different days
                if aktivitet1 == "Kino i hall 2":
                    # NB! Elever i samme klasse, vil kanskje på samme dag?? (ikke implementert - det kan kræsje med andre aktiviteter)
                    if random.randint(0,1) == 1:
                        aktivitetsdag = "mandag"
                    else:
                        aktivitetsdag = "tirsdag"
                else:
                    aktivitetsdag = fordelingDag[aktivitet1]
                # Appending application
                # NB: Ønsker å flytte prioritet bakerst, for å forenkle koden for oppsett i Word.
                allApplications.append((tup[inputTupleNamePosition], tup[inputTupleClassPosition], aktivitet1, highPriorityString, aktivitetsdag)) # navn == [9], klassse == [14], aktivitet1 == [17] 
                # (navn, klasse, aktivitet, prioritet, dag)
                #print((tup[name], tup[studentClass], aktivitet1, highPriorityString))
    print("Tester dag for aktivitet", allApplications[0])
    #print(f"Debug allApplicaitions: len(application): {len(allApplications[0])} allApplications[0]: {allApplications[0]} len(allApplications()): {len(allApplications)}")
    
    return allApplications


# All applications are sorted into groups according to priority (Ønsker å endre programmet til numerisk prioritet?)
# Pulling tuples of all rows from the Excel file
allApplications = [tuple(x) for x in df.itertuples(index=False, name=None)]

# Checking if a student have submittet more than one response
def find_duplicates(applicationTuples):
    seen = set()
    duplicates = []
    for item in applicationTuples:
        if item[inputTupleNamePosition].lower() in seen:
            duplicates.append(item[inputTupleNamePosition])
        else:
            seen.add(item[inputTupleNamePosition].lower())
    return duplicates

# Finding and printing students that have submitted more than one reply
# vil kanskje integrere dette i structure_applications()
duplicates = find_duplicates(allApplications)
for student in duplicates:
    print("DUPLICATE STUDENT: ", student)  

# Restructuring applications, creating one application for each desired activity
print("Restructuring allApplications")
allApplications = structure_applications(allApplications)
finalAssembly = allApplications[:] # Lager en kopi av allApplications til dokumentene til slutt i prosessen.
random.shuffle(allApplications) # Trenger å randomisere for rettferdig fordeling (riktig sted?)

gruppertTotal = [] # Er det bare gruppertTotal som brukes?
# Prioritet skal implementeres som verdi/tall
#denne kan umulig passe å definere her. Skal den defineres på starten? Skal den være en dict?

# Dict to keep track of what days students have been assigned activities
    # -> Bør den kanskje være en list-of-tuples, for å pare aktivitet og dag?
eleverMedBekreftedeAktiviteter = {elev: [] for elev in allePaameldteElever}

print("Setup complete")

# priorityString skal erstattes med numerisk prioritet
def group_student_applications(priorityString): 
    # selecting applications corresponding to selected priority for the application
    groupOfPrioritizedApplications = None
    # Returning group of applications with requested priority
    if priorityString: # will turn into an integer
        groupOfPrioritizedApplications = [tuple for tuple in allApplications if tuple[applicationTuplePriorityPosition] == priorityString]
    else: 
        groupOfPrioritizedApplications = [tuple for tuple in allApplications]
    # creating groups for number of applications per student
    groups = {key: [] for key in range(1, maxApplicationsPerStudent)} # 0 applications should not be possible -> Yet it is, when applicatiions are removed due to full activities
    # counting numbre of applications for each student (filtered by priority)
    counter = Counter(application[applicationTupleNamePosition] for application in groupOfPrioritizedApplications) 

    # adding previously assigned activities to number of applications (if not fully assigned)
    # NB! Denne må muligens endres for å ta hensyn til hvilken dag aktiviteten er lagt til.  
    for student in eleverMedBekreftedeAktiviteter: 
        if len(eleverMedBekreftedeAktiviteter[student]) < maxActivitiesPerStudent and student in counter:
            counter[student] += len(eleverMedBekreftedeAktiviteter[student])
    for student, numberOfApplications in counter.items():
        groups[numberOfApplications].append(student)
        
    # returning group of students sorted by numberOfApplicationsPerStudent, filtered by priority of applications    
    return groups 

def place_student(studentName):
    # NB - VI MÅ HA LOGIKK FOR Å IKKE BOOKE TO AKTIVITETER PÅ SAMME DAG! (fikset?)
    for application in list(allApplications):
        if application[applicationTupleNamePosition] == studentName:

            ## Bug! Det er noe feil med dagene (fikset?)

            # Hack to place students applying for "Python for de som kan litt fra før" first. Delete or comment out if functionality not desired.
            if studentName in listOfStudentsWithPythonApplications:
                if application[applicationTupleActivityPosition] == "Python for de som kan litt fra før":
                    listOfStudentsWithPythonApplications.remove(studentName)
                    fordeling[application[applicationTupleActivityPosition]].append(studentName) # Adding student to group
                    allApplications.remove(application)
                    eleverMedBekreftedeAktiviteter[studentName].append(fordelingDag[application[applicationTupleActivityPosition]]) 
                    return allApplications, fordeling, eleverMedBekreftedeAktiviteter, allePaameldteElever
                else:
                    continue

            # Removing application and keep searching if activity is full
            if fordelingMax[application[applicationTupleActivityPosition]] == len(fordeling[application[applicationTupleActivityPosition]]):
                allApplications.remove(application) #FUNKER
                continue # Sjekk at continue funker som forutsatt - FUNKER (nesten sikker)
            
            # If student is assigned to max number of activities, application is removed
            if len(eleverMedBekreftedeAktiviteter[studentName]) == maxActivitiesPerStudent:
                # VIKTIG - SJEKK FOR AT TO AKTIVITETER IKKE BOOKES PÅ SAMME DAG (fikset?)
                allApplications.remove(application)
                return allApplications, fordeling, eleverMedBekreftedeAktiviteter, allePaameldteElever
            
            # If student is already booked this day, application is removed, and algorithm keeps searching
            if fordelingDag[application[applicationTupleActivityPosition]] in eleverMedBekreftedeAktiviteter[studentName]:
                allApplications.remove(application)
                continue
            allApplications.remove(application)
            fordeling[application[applicationTupleActivityPosition]].append(studentName) # Adding student to group

            # adding activityDay to list of students
            eleverMedBekreftedeAktiviteter[studentName].append(fordelingDag[application[applicationTupleActivityPosition]]) 
            return allApplications, fordeling, eleverMedBekreftedeAktiviteter, allePaameldteElever
    # In case activity is fully signed for this student, and the student has no further applications
    return allApplications, fordeling, eleverMedBekreftedeAktiviteter, allePaameldteElever

# DEBUG / hack
listOfStudentsWithPythonApplications = []
for app in allApplications:
    if app[applicationTupleActivityPosition] == "Python for de som kan litt fra før":
        listOfStudentsWithPythonApplications.append(app[applicationTupleNamePosition])
# print(fordeling["Python for de som kan litt fra før"])
    

# Skriv om til function(s)
while allApplications:

    # While there are still students to assign that have (applications + assigned) <= maxActivities
    gruppertTotal = group_student_applications("") #Sending empty priority string
    #def group_student_applications(priorityString, allApplications, fordeling, maxActivitiesPerStudent, maxApplicationsPerStudent): 
    #print("\nAssigning students to activities\n")
    exitCounter2 = 0 #fjernes
    # Defining groups that have numberOfApplications <= maxActivitiesPerStudent
    priorityGroupKeys = [number for number in range(1, maxActivitiesPerStudent + 1)]
    # Legg inn sperre-betingelse for å ikke kjøre den unødig
    # while gruppertTotal(few) != eleverMedFordelteAktiviteter

    # check for redundancy
    runWhileUnassignedRemains = True
    #while not all(len(gruppertTotal[group]) == 0 for group in priorityGroupKeys): # "group in gruppertTotal and" kan legges til for mer robust kode (hindrer KeyError)
    # Denne while-setningen feiler fordi assigned legges til gruppert.
    fewApplicationsLoopCounter = 0
    # listOfStudentsWithFewApplications har bare gyldighet inni loopen. 
    while runWhileUnassignedRemains: # Hva gjør runWhileUnnassigned nå? == runWhileTrue ??
        fewApplicationsLoopCounter += 1
        # DEBUG
        #print("Assigning students with few applications: ", fewApplicationsLoopCounter)
        # VIKTIG: Er denne feil? Var den feil? Er den feil nå?
        eleverFullyAssigned = [student for student in eleverMedBekreftedeAktiviteter if len(eleverMedBekreftedeAktiviteter) == maxActivitiesPerStudent]
        # expression for key in dict if condition ==> Bør være len(eleverMedBekreftedeAktiviteter[student]) ??
#        print("Debug - eleverFullyAssigned: ", eleverFullyAssigned)
#        print("Debug - len(eleverMedBekreftedeAktiviteter): ", len(eleverMedBekreftedeAktiviteter))
        #print("Elever fully assigned:\n", eleverFullyAssigned)
        # Elever som er assigned vil være med her. 
        # Elever som er fully assigned, vil være fjernet (? sjekk i group_student_applications!)
        # Vil ha en list-of-lists får å plassere elever med færrest ønsker først.
        listOfStudentsWithFewApplications = [gruppertTotal[gruppe] for gruppe in priorityGroupKeys] # maxActivitiesPerStudent lister med elever 
        # Funker denne som forutsatt? Tviler!
        testliste = [student for listen in listOfStudentsWithFewApplications for student in listen]
        #if all(student in eleverFullyAssigned for student in listOfStudentsWithFewApplications):
        if all(student in eleverFullyAssigned for student in testliste): # Prøver ny logikk (eventuelt: pakk ut listOfStudentsWithFewApplications)
        # (condition for element in iterable)
        # Funker nå?
            # Tror ikke denne funker - prøver ny logikk
                # -> Tror den trigges for ofte
            #print("\n\nNo more students with few applications - break\n\n")
            break # Bør heller manipulere runWhileUnassignedRemains? ->
        # breaks this loop after a few itereations -> should be -> if no student needs priority placement
            # Then goes to random assignment for remaining students
            #redundant nå?
        if exitCounter2 == 30:
            break
        exitCounter2 += 1
#        print("Debug exitCounter: ", debugExitCounter)
        # CODE
        # listOfStudentsWithFewApplications har nå alle elevene med 1 og 2 applications/assignments. 
        # de kan plasseres i fordelt
        debugBreakCounter = 0
        repeatFew = True
        while repeatFew: # Gjør denne det samme som løkken over? eller trenger den ekstra iterasjoner?
        # NB! list-of-list er truthy selv om den et tom!
        # listOfStudentsWithFewApplications inkluderer elever som er assigned (gjelder det også fully-assigned?)
        # dvs. den blir aldri False
            debugBreakCounter += 1
            if debugBreakCounter == 4:
                break
            # print("allApplications: ", len(allApplications))
            # for i in gruppertTotal:
            #     if len(gruppertTotal[i]) > 0:
            #         print(f"gruppertTotal[{i}] {len(gruppertTotal[i])}")
            listIndex = 0
            # Making sure that students with fewer applications are assigned first.
            # Lurer på om listOfStudentsWithFewApplications kaller place_student unødig.
            #print("Tester at listen er 2 lang:\t", len(listOfStudentsWithFewApplications))
            for i in range(len(listOfStudentsWithFewApplications)): #redundant?
                if listOfStudentsWithFewApplications[i]:
                    listIndex = i # Usikker på denne logikken; skal det være +1? -> Tror ikke det
                # Placing student in activity-group according to application
                #print(f"Length of listOfStudentsWithFewApplications[{listIndex}]: ", len(listOfStudentsWithFewApplications[listIndex]))
                for student in listOfStudentsWithFewApplications[listIndex]:
                    #print("PLACING STUDENT LOOP")
                    #print("fewApplications", len(listOfStudentsWithFewApplications[0]), len(listOfStudentsWithFewApplications[1]))
                    # Trenger vel ikke returnere allePaamledteElever?
                    #print("len(fordeling)1: ", sum(len(lst) for lst in fordeling.values()))
                    allApplications, fordeling, eleverMedBekreftedeAktiviteter, allePaameldteElever = place_student(student)
                    #print("len(fordeling)2: ", sum(len(lst) for lst in fordeling.values()))
                    # Kontroller at listen regenereres korrekt, spesielt at gruppertTotal[] er oppdatert, og at priorityGroupKeys er korrekt
                    # pass på at variablene eksisterer i riktig form til å sendes til funksjonen.
            # Students that are fully assigned, no longer have applications in allApplications (some unassigned students may have been removed)
            # Recreating groups of students with number of applications as keys
            gruppertTotal = group_student_applications("") # Sending empty priority string
            listOfStudentsWithFewApplications = [gruppertTotal[gruppe] for gruppe in priorityGroupKeys]
            # Checking if there still are unassigned applications for listOfStudentsWithFewApplications
                # (listOfStudentsWithFewApplications also include students with assigned activities)
            # Kontroller til slutt at denne koden ikke er redundant.
            repeatFew = False
            for i in range(len(listOfStudentsWithFewApplications)):
                for student in list(listOfStudentsWithFewApplications[i]):
                    for application in allApplications:
                        if application[applicationTupleNamePosition] == student:
                            repeatFew = True # kjører løkka om igjen om noen applications ikke er plassert
                    if repeatFew:
                        break
                if repeatFew:
                    break
        # end-while listOfStudentsWithFewApplications
    # end-while runWhileUnassignedRemains - ERSTATTES av bedre logikk
    
    # Itererer over keys (fordi det er en dict) og lager liste av elever som ikke er fully assigned
    # when no student needs priority assignment, this code assignes students in random order (ikke implementert?)
    # NB! elever kan ha 0-1 bekreftede aktiviteter, men ingen flere applications!
    # allApplications should now contain only applications for students that have more applications than needed to be fully assigned
    setOfUnassignedStudents = set([application[applicationTupleNamePosition] for application in allApplications])
    setOfStudentsWithHighPriorityApplications = set([application[applicationTupleNamePosition] for application in allApplications if application[applicationTuplePriorityPosition] == highPriorityString])
    #print("Assigning students without fewApplications")
    # Bruker jeg ikke place_student her??
    for student in setOfUnassignedStudents: 
        highPriorityApplicationExists = False
        fullyAssignedGroupBreak = False
        if student in setOfStudentsWithHighPriorityApplications:
            highPriorityApplicationExists = True
        for application in list(allApplications):
            if application[applicationTupleNamePosition] == student:
                # Prioritizing high-priority applications
                if application[applicationTuplePriorityPosition] == highPriorityString or not highPriorityApplicationExists:
                    # Don't need to check for fully-assigned here
                    # IMPORTANT! return value needed? - Not part of a function (yet), so: NO
                    #print("Placing student with several applications", len(allApplications))
                    allApplications, fordeling, eleverMedBekreftedeAktiviteter, allePaameldteElever = place_student(student)
                    if len(fordeling[application[applicationTupleActivityPosition]]) == fordelingMax[application[applicationTupleActivityPosition]]:
                        fullyAssignedGroupBreak = True 
                        print("Group full: ", application[applicationTupleActivityPosition]) 
                        break
        # When a group is full, we need to clean applications and re-ckeckfor few allpications
        if fullyAssignedGroupBreak: 
            break # for student in listOfUnnassignedStudents
    #print("Number of remaining applications: ", len(allApplications))
# end while allApplications
                
#print("DEBUG EXIT_COUNTER: ", debugExitCounter)


# Creating dict with students that aren't fully assigned
studentsWithUnassignedActivities = {i:[] for i in range(1, maxActivitiesPerStudent + 1)}
#print("unassigned: ", studentsWithUnassignedActivities)
#DEBUG ==> skal skrive ut egen liste / legges til klassene
for student in eleverMedBekreftedeAktiviteter:
    for numberOfActivities in range(0, maxActivitiesPerStudent):
        if eleverMedBekreftedeAktiviteter[student] == numberOfActivities:
            studentsWithUnassignedActivities[maxActivitiesPerStudent - numberOfActivities].append(student)
print("studentsWithUnassignedActivities: ", len(studentsWithUnassignedActivities), studentsWithUnassignedActivities)

# Hver elev skal få oppfylt 2-to ønsker. 
# Elevene med 1-2 ønsker, må få sine oppfylt først.
# Det må kontrolleres om elever med 3-4 ønsker har fått sine redusert
# Deretter må resten få oppfylt sitt første ønske
# Ny kontroll må gjennomføres for hver gruppe som blir full
# Til slutt må alle gjenværende elever få oppfylt sitt andre ønske, med prioritet til "Vil mye"
# Advarsel må skrives for hver elev som ikke får oppfylt 2 ønsker

#WARNING: Systemet kan games; ved å oppgi få ønsker, er det større sjanse for å få akkurat hva du ønsker deg.
    # Elever kan risikere å ikke få en aktivitet om ønskede aktiviteter er fyllt opp.
# Hvis det blir fullt, kan elever settes opp på en tilfeldig aktivitet?



# CREATING DOCUMENTS:

#def create_word_document(file_name, headline, dictOfTuples):
def create_word_document(dictOfTuples):

    # Add data (try-except)
    for group in dictOfTuples:
        print(group) 
        file_name = group

        # Create a new Document
        doc = Document()

        # Set paper size to A4 (redundant?)
        sections = doc.sections
        for section in sections:
            section.page_width = Pt(595.3)
            section.page_height = Pt(841.9)
            section.left_margin = Pt(50)
            section.right_margin = Pt(50)
            section.top_margin = Pt(0)
            section.bottom_margin = Pt(30)

        # Add headline
        h = doc.add_heading(level=1)
        run = h.add_run(group.capitalize())
        run.font.name = 'Calibri Light'
        run.font.size = Pt(14)
        run.bold = True

        h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a table
        kolonner = 5 # I tilfelle listen er tom (skal være 5)
#        kolonner = len(dictOfTuples[group][0]) + 1
        table = doc.add_table(rows=1, cols=kolonner) 

        # Style the table
        table.style = 'Table Grid'

        # Add the header row
        hdr_cells = table.rows[0].cells
        overskrifter = ["Navn", "Klasse", "Aktivitet", "Dag"]
        for i, kolonne in enumerate(overskrifter):
            run = hdr_cells[i + 1].paragraphs[0].add_run(kolonne)
            run.font.name = 'Calibri Light'
            run.font.size = Pt(10)
            run.bold = True 

        # Adding data to table
        for x, item in enumerate(dictOfTuples[group]): # Sjekk om indeksene må justeres for header
            row_cells = table.add_row().cells
            # Numbering the applications
            run = row_cells[0].paragraphs[0].add_run(str(x + 1))
            # Adding data from applications
            for index, item2 in enumerate(item): 
                # Hack to omit priority column (should switch places with 'Day' in later iterations)
                if index >= 3:
                    adjustedIndex = index - 1
                else:
                    adjustedIndex = index
                if index != 3:
                    run = row_cells[adjustedIndex +1].paragraphs[0].add_run(item2) # Har en ekstra kolonne nå; fjerne prioritet(?)
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)
                run.bold = False


        # Set different column widths
        column_widths = [20, 150, 60, 250]
        for col_index, width in enumerate(column_widths):
            for cell in table.columns[col_index].cells:
                cell.width = Pt(width)

        # Save the document
        doc.save(f"{file_name}.docx")

# VALIDATING (ting funker ikke - finn ut hvorfor)
def validation(allePaameldteElever, eleverMedBekreftedeAktiviteter):
    # eleverMedBekreftedeAktiviteter skal ha oversikten over alt
    # fordeling skal matche med eleverMed...
    # studentsWithUnassignedActivities skal matche med eleverMed... med aktiviteter < 2
    # Det må sjekkes for om elevenes aktiviteter er på to forskjellige dager

    # Checking if all students have been accounted for
    for elev in allePaameldteElever:
        if elev not in eleverMedBekreftedeAktiviteter:
            raise Exception(f"Student {elev} has not been assigned to activities") 
    print("CONTROL COMPLETED: All students accounted for")

    # Checking if all students are assigned to maxActivitiesPerStudent
    unassigned_counter = 0
    dictOfUnassigned = {"unassigned": []}
    # dictOfActivities = {activity: [] for activity in fordeling}
    # setOfClasses = {tup[1] for tup in finalAssembly}
    # dictOfClasses = {klasse: [] for klasse in setOfClasses}
    #dictOfClasses[application[applicationTupleClassPosition]].append(application)
    #dictOfClasses["unassigned"].append(application)
    for elev in eleverMedBekreftedeAktiviteter:
        if len(eleverMedBekreftedeAktiviteter[elev]) < maxActivitiesPerStudent:
            print(f"WARNING: student {elev} is not fully assigned")
            #print(eleverMedBekreftedeAktiviteter[elev])
            unassigned_counter += 1
            for application in finalAssembly:
                if application[applicationTupleNamePosition] == elev:
                    print(fordelingDag[application[applicationTupleActivityPosition]])
                    dictOfUnassigned["unassigned"].append(application)
    print("CONTROL COMPLETED")
    print(f"WARNING: {unassigned_counter} students have not been fully assigned")

    dictOfUnassigned["unassigned"] = sorted(dictOfUnassigned["unassigned"], key=lambda x: x[applicationTupleNamePosition])

    return dictOfUnassigned
    
dictOfUnassigned = validation(allePaameldteElever, eleverMedBekreftedeAktiviteter)


# VIKTIG - Sjekk/tell kollisjon med dager
# hvilke aktivitetspar er mest utsatt?

# CREATING WORD-DOCUMENTS

# Preparing data for creation of Word-documents
def prepare_data_for_word_documents():
    dictOfActivities = {activity: [] for activity in fordeling}
    setOfClasses = {tup[1] for tup in finalAssembly}
    dictOfClasses = {klasse: [] for klasse in setOfClasses}
    for activity in fordeling:
        for student in fordeling[activity]:
            for application in finalAssembly:
                # print(application) # Funker
                if student == application[applicationTupleNamePosition] and activity == application[applicationTupleActivityPosition]:  
                    dictOfActivities[activity].append(application)
                    dictOfClasses[application[applicationTupleClassPosition]].append(application)
                    #print(application)

    # Sort the lists by name within the dictionaries
    for activity in dictOfActivities:
        dictOfActivities[activity] = sorted(dictOfActivities[activity], key=lambda x: x[applicationTupleNamePosition])

    for klasse in dictOfClasses:
        dictOfClasses[klasse] = sorted(dictOfClasses[klasse], key=lambda x: x[applicationTupleNamePosition])

    return dictOfActivities, dictOfClasses

# Presenting data as Word documents
dictOfActivities, dictOfClasses = prepare_data_for_word_documents()
create_word_document(dictOfActivities)
create_word_document(dictOfClasses)
create_word_document(dictOfUnassigned)
# NB! "Kino i Hall 2" kan spille algoritmen et puss. Det ser ut som "fordelt" ikke lar den ligge på 'tirsdag' (for den enkelte elev)
# En elev er satt op med mandag-mandag i 'validation', men har mandag-tirsdag i unnassigned.docx
# Mulig "Kino i Hall 2" skal fordeles til slutt. Kanskje også til elever som mangler 1 aktivitet??)
# Trenger den kanskje ikke fikses på? Fikses ved bedre design av spørreskjemaet?
