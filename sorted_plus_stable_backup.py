#from curses import doupdate # Hvor kommer denne fra??
import pandas as pd
from collections import Counter
import random

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# def __init__

# NB! Skal all input gjøres til .lower() ?

# target Excelark bør være userInput

df = pd.read_excel('Påmelding til turer!.xlsx')
fulltNavn = df['Skriv inn fornavn og etternavn'].tolist() # unique Identifier
klasse = df['Hvilken klasse går du i?'].tolist() # Class
Aktivitet1 = df['Hvilke aktiviteter vil du være med på? Velg så mange du vil, så vil du bli satt opp på to aktiviteter.'].tolist() # Acting as wanted group
# aktivitet2 = df['Hvilken aktivitet vil du være med på tirsdag 11. juni. Velg en ny aktivitet, ikke den du valgte for mandag.'].tolist() # Acting as wanted group
#priority = df['Prioritet'] # Wishes are sorted by priority
#friends = [] # Used to group friends together

# Flere verdier kan være userInput -> GUI vurderes etter hvert
highPriorityString = "Dette har jeg veldig lyst til"
lowPriorityString = "Dette har jeg litt lyst til"
maxActivitiesPerStudent = 2 # Bør være userInput
# Bør hentes fra Excelark
maxApplicationsPerStudent = 14 # aka totalNumberOfActivities
# All available unique students
allePaameldteElever = set(fulltNavn) # Kanskje heller set(df['Fullt navn'].tolist())
#VIKTIG! første oppføring er 'nan'
#print("DEBUG: allePaameldteElever ", allePaameldteElever)

# Pass på at index til aktivitetene er riktig med ny layout på Forms
def structure_applications(inputApplications):
    # print(inputApplications[0])
    # print(inputApplications[0][9], inputApplications[0][14], inputApplications[0][17])
    # Stemmer
    allApplications = []
    for tup in inputApplications:
        for aktivitet1 in tup[17].split(';'): # 9 er navn 14 er klasse, 17 er aktivitet
            if aktivitet1: # Fixing trailing ';' (perhaps)
                allApplications.append((tup[9], tup[14], aktivitet1, highPriorityString)) # navn == [9], klassse == [14], aktivitet1 == [17] 
                #print((tup[8], tup[11], aktivitet1, highPriorityString))
        # for aktivitet2 in tup[17].split(';'):
        #     if aktivitet2:
        #         allApplications.append((tup[8], tup[11], aktivitet2, highPriorityString))
    #print(f"Debug allApplicaitions: len(application): {len(allApplications[0])} allApplications[0]: {allApplications[0]} len(allApplications()): {len(allApplications)}")
    
    return allApplications
# husk å modifisere 'fordeling' og 'fordelingMax'


# All applications are sorted into groups according to priority (kan programmet endres til numerisk prioritet?)
# kanskje flytte Hoy og Lav til def gruppert
# Prøver å unngå å bruke dem. Redundant?
#prioritetHoy = [tuple(x) for x in df.itertuples(index=False, name=None) if x[3] == highPriorityString]
#prioritetLav = [tuple(x) for x in df.itertuples(index=False, name=None) if x[3] == lowPriorityString]
# Pulling tuples of all rows from the Excel file
allApplications = [tuple(x) for x in df.itertuples(index=False, name=None)]
# Rstructuring appliations, creating one application for each desired activity
print("Restructuring allApplications")
#DEBUG/TEST:
pythonCounter = 0
for app in allApplications:
    if "kan litt fra før" in app[17]:
        pythonCounter += 1
print("Viderekomne: ", pythonCounter)
print("INGEN GRUPPER HAR BEGRENSING => SKAL VI PRIORITERE GRUPPER MED FÅ ELLER MANGE PÅMELDINGER?")
allApplications = structure_applications(allApplications)
# for app in allApplications:
#     print(app)
finalAssembly = allApplications[:] # Lager en kopi av allApplications til dokumentene til slutt i prosessen.
random.shuffle(allApplications) # Trenger å randomisere for rettferdig fordeling (riktig sted?)
#Removing redundant applications (midlertidig deaktivert for testing purposes.)
#allApplications = list(set(allApplications)) # NOT TESTED (reduserer antall applications til 140?)
# priority groups; built in group_student_applications() 
# Skal den kalles her?

#gruppertHoy = [] # Er disse redundante? muligens lurt å ha dem definert
#gruppertLav = []
gruppertTotal = [] # Er det bare gruppertTotal som brukes?

# Creating activities with max number of participants
# NB! Bør hentes fra excel/dataframe -> Komplisert, valgene er ikke tilgjengelige som en enhet fra Excel.
#fordeling = {"emel": [], "anne marie": [], "sveinung": [], "natasha": [], "elisabeth": [], "andreas": []}
fordeling = {"Orientering": [], "Frisbee Ultimate": [], "Kino i hall 2": [], \
    "Langtur i Østmarka": [], "Python nybegynner": [], "Brettspilldag": [], \
        "Ta med bok og les - det servers drikke!": [], "Python for de som kan litt fra før": [], \
            "Trigonometri": [], "Språkkafé": [], "Tur i Botanisk Hage + utendørsklatring i buldreveggen på Enerhaugen": [], \
                "Basket i Kubaparken": [], "Joggetur rundt Maridalsvannet": [], "Piknik i Torshovparken": []}
fordelingMax = {"Orientering": 100, "Frisbee Ultimate": 100, "Kino i hall 2": 100, \
    "Langtur i Østmarka": 100, "Python nybegynner": 100, "Brettspilldag": 100, \
        "Ta med bok og les - det servers drikke!": 100, "Python for de som kan litt fra før": 100, \
            "Trigonometri": 100, "Språkkafé": 100, "Tur i Botanisk Hage + utendørsklatring i buldreveggen på Enerhaugen": 100, \
                "Basket i Kubaparken": 100, "Joggetur rundt Maridalsvannet": 100, "Piknik i Torshovparken": 100}
fordelingDag = {"Orientering": "mandag", "Frisbee Ultimate": "mandag", "Kino i hall 2": "mandag", \
    "Langtur i Østmarka": "mandag", "Python nybegynner": "mandag", "Brettspilldag": "tirsdag", \
        "Ta med bok og les - det servers drikke!": "tirsdag", "Python for de som kan litt fra før": "tirsdag", \
            "Trigonometri": "tirsdag",  "Språkkafé": "tirsdag", "Tur i Botanisk Hage + utendørsklatring i buldreveggen på Enerhaugen": [], \
                "Basket i Kubaparken": "tirsdag", "Joggetur rundt Maridalsvannet": "tirsdag", "Piknik i Torshovparken": "tirsdag"}
# VIKTIG - bør kanskje fordele aktiviteter på dager? -> DONE!
# NB! Bør være userInput
#fordelingMax = {"emel": 120, "anne marie": 30, "sveinung": 25, "natasha": 50, "elisabeth": 8, "andreas": 100}
#fordelingMax = {"Orientering": 100, "Frisbee Ultimate": 100, "Kino i hall 2": 100, "Langtur i Østmarka": 100, "Python nybegynner": 100, "Brettspilldag": 100}
# Keeping track of number of activities each student is assigned to
# kanskje allePaamledteElever er redundant ? 

eleverMedBekreftedeAktiviteter = {elev: [] for elev in allePaameldteElever}

print("Setup complete")



# slått sammen med count_applications (funker det?)
# 'fordeling' brukes ikke? Er flere variabler endret og skal returneres?
def group_student_applications(priorityString, allApplications, fordeling, maxActivitiesPerStudent, maxApplicationsPerStudent): 
    #print("\tGrouping student applications function ")
    # selecting applications corresponding to selected priority for the application
    groupOfPrioritizedApplication = None
    if priorityString:
        groupOfPrioritizedApplication = [tuple for tuple in allApplications if tuple[3] == priorityString]
    else: 
        groupOfPrioritizedApplication = [tuple for tuple in allApplications]
# Trying to merge functions
    # Redundant -> kommentert ut
    #gruppert = count_applications(groupOfPrioritizedApplication) # 'groupOfPrioritizedApplication' er list-of-tuples(applications) med ønsket prioritet
    # creating groups for number of applications per student
    groups = {key: [] for key in range(1, maxApplicationsPerStudent)} # 0 applications should not be possible -> Yet it is, when applicatiions are removed due to full activities
    # counting applications per student (filtered by priority)
    counter = Counter(application[0] for application in groupOfPrioritizedApplication) # app[0] er elevnavn; teller hvor mange applications eleven har sendt

    # adding assigned activities to number of applications (if not fully assigned)
#    print("counter: ", counter)
    # navn: numberOfApplications
    for student in eleverMedBekreftedeAktiviteter: 
        if len(eleverMedBekreftedeAktiviteter[student]) < maxActivitiesPerStudent and student in counter:
            counter[student] += len(eleverMedBekreftedeAktiviteter[student])
    #print(counter)
    print("DEBUG: line 130") # 12 not found
    for student, numberOfApplications in counter.items():
        groups[numberOfApplications].append(student)
    # returning group of student sorted by numberOfApplicationsPerStudent, filteres by priority of applications    
    return groups # 'gruppert' er dict med elever fortelt etter antall applications. 
    # Skal activitiesAsssigned legges til gruppene?? Tror det. 

# Funksjonen bør returnere nye verdier
def place_student(allApplications, fordeling, studentName, eleverMedBekreftedeAktiviteter, allePaameldteElever):
    # NB - VI MÅ HA LOGIKK FOR Å IKKE BOOKE TO AKTIVITETER PÅ SAMME DAG!
    for application in list(allApplications):
        if application[0] == studentName:
            # Removing application and keep searching if activity is full
            if fordelingMax[application[2]] == len(fordeling[application[2]]):
                allApplications.remove(application) #FUNKER
                continue # Sjekk at continue funker som forutsatt - FUNKER (nesten sikker)
            # If student is assigned to max number of activities, application is removed
            if len(eleverMedBekreftedeAktiviteter[studentName]) == maxActivitiesPerStudent:
                # VIKTIG - SJEKK FOR AT TO AKTIVITETER IKKE BOOKES PÅ SAMME DAG
                allApplications.remove(application)
                return allApplications, fordeling, eleverMedBekreftedeAktiviteter, allePaameldteElever
            # If student is already booked this day, applications is removed, and algorithm keeps searching
            if fordelingDag[application[2]] in eleverMedBekreftedeAktiviteter[studentName]:
                # print("FINDING APPLICATION FOR ANOTHER DAY")
                # print("Debug: fordelingDag - ", fordelingDag[application[2]],"gruppe: ", application[2], "bekreftedeAktiviteter: ", eleverMedBekreftedeAktiviteter[studentName])
                allApplications.remove(application)
                continue
            allApplications.remove(application)
            # VIKTIG - TEST OM DAGEN FOR AKTIVITET ER BOOKET.
            fordeling[application[2]].append(studentName) # Adding student to group
            eleverMedBekreftedeAktiviteter[studentName].append(fordelingDag[application[2]]) 
#            print("bookede dager: ", eleverMedBekreftedeAktiviteter[studentName])
            return allApplications, fordeling, eleverMedBekreftedeAktiviteter, allePaameldteElever
    # In case activity is fully signed for this student, and the student has no further applications
    return allApplications, fordeling, eleverMedBekreftedeAktiviteter, allePaameldteElever
applicationCounter = 1 #brukes denne til noe?
# debugExitCounter = 0 # DEBUG
# Skriv om til function(s)
while allApplications: # trenger continue(?)

    # While there are still students to assign that have (applications + assigned) <= maxActivities
    gruppertTotal = group_student_applications("", allApplications, fordeling, maxActivitiesPerStudent, maxApplicationsPerStudent)
    #def group_student_applications(priorityString, allApplications, fordeling, maxActivitiesPerStudent, maxApplicationsPerStudent): 
    #print("\nAssigning students to activities\n")
    exitCounter2 = 0
    # Defining groups that have numberOfApplications <= maxActivitiesPerStudent
    priorityGroupKeys = [number for number in range(1, maxActivitiesPerStudent + 1)]
    # Legg inn sperre-betingelse for å ikke kjøre den unødig
    # while gruppertTotal(few) != eleverMedFordelteAktiviteter

    # check for redundancy
    runWhileUnassignedRemains = True
    #while not all(len(gruppertTotal[group]) == 0 for group in priorityGroupKeys): # "group in gruppertTotal and" kan legges til for mer robust kode (hindrer KeyError)
    # Denne while-setningen feiler fordi assigned legges til gruppert.
    fewApplicationsLoopCounter = 0
    # listOfStudentsWithFewApplications har bare gyldighet inni loopen. 
    while runWhileUnassignedRemains: # Hva gjør runWhileUnnassigned nå? == runWhileTrue ??
        fewApplicationsLoopCounter += 1
        # DEBUG
        #print("Assigning students with few applications: ", fewApplicationsLoopCounter)
        # VIKTIG: Er denne feil? Var den feil? Er den feil nå?
        eleverFullyAssigned = [student for student in eleverMedBekreftedeAktiviteter if len(eleverMedBekreftedeAktiviteter) == maxActivitiesPerStudent]
        # expression for key in dict if condition ==> Bør være len(eleverMedBekreftedeAktiviteter[student]) ??
#        print("Debug - eleverFullyAssigned: ", eleverFullyAssigned)
#        print("Debug - len(eleverMedBekreftedeAktiviteter): ", len(eleverMedBekreftedeAktiviteter))
        #print("Elever fully assigned:\n", eleverFullyAssigned)
        # Elever som er assigned vil være med her. 
        # Elever som er fully assigned, vil være fjernet (? sjekk i group_student_applications!)
        # Vil ha en list-of-lists får å plassere elever med færrest ønsker først.
        listOfStudentsWithFewApplications = [gruppertTotal[gruppe] for gruppe in priorityGroupKeys] # maxActivitiesPerStudent lister med elever 
        # Funker denne som forutsatt? Tviler!
        testliste = [student for listen in listOfStudentsWithFewApplications for student in listen]
        #if all(student in eleverFullyAssigned for student in listOfStudentsWithFewApplications):
        if all(student in eleverFullyAssigned for student in testliste): # Prøver ny logikk (eventuelt: pakk ut listOfStudentsWithFewApplications)
        # (condition for element in iterable)
        # Funker nå?
            # Tror ikke denne funker - prøver ny logikk
                # -> Tror den trigges for ofte
            #print("\n\nNo more students with few applications - break\n\n")
            break # Bør heller manipulere runWhileUnassignedRemains? ->
        # breaks this loop after a few itereations -> should be -> if no student needs priority placement
            # Then goes to random assignment for remaining students
            #redundant nå?
        if exitCounter2 == 30:
            break
        exitCounter2 += 1
#        print("Debug exitCounter: ", debugExitCounter)
        # CODE
        # listOfStudentsWithFewApplications har nå alle elevene med 1 og 2 applications/assignments. 
        # de kan plasseres i fordelt
        debugBreakCounter = 0
        repeatFew = True
        while repeatFew: # Gjør denne det samme som løkken over? eller trenger den ekstra iterasjoner?
        # NB! list-of-list er truthy selv om den et tom!
        # listOfStudentsWithFewApplications inkluderer elever som er assigned (gjelder det også fully-assigned?)
        # dvs. den blir aldri False
            debugBreakCounter += 1
            if debugBreakCounter == 4:
                break
            # print("allApplications: ", len(allApplications))
            # for i in gruppertTotal:
            #     if len(gruppertTotal[i]) > 0:
            #         print(f"gruppertTotal[{i}] {len(gruppertTotal[i])}")
            listIndex = 0
            # Making sure that students with fewer applications are assigned first.
            # Lurer på om listOfStudentsWithFewApplications kaller place_student unødig.
            #print("Tester at listen er 2 lang:\t", len(listOfStudentsWithFewApplications))
            for i in range(len(listOfStudentsWithFewApplications)): #redundant?
                if listOfStudentsWithFewApplications[i]:
                    listIndex = i # Usikker på denne logikken; skal det være +1? -> Tror ikke det
                # Placing student in activity-group according to application
                #print(f"Length of listOfStudentsWithFewApplications[{listIndex}]: ", len(listOfStudentsWithFewApplications[listIndex]))
                for student in listOfStudentsWithFewApplications[listIndex]:
                    #print("PLACING STUDENT LOOP")
                    #print("fewApplications", len(listOfStudentsWithFewApplications[0]), len(listOfStudentsWithFewApplications[1]))
                    # Trenger vel ikke returnere allePaamledteElever?
                    #print("len(fordeling)1: ", sum(len(lst) for lst in fordeling.values()))
                    allApplications, fordeling, eleverMedBekreftedeAktiviteter, allePaameldteElever = place_student(allApplications, fordeling, student, eleverMedBekreftedeAktiviteter, allePaameldteElever)
                    #print("len(fordeling)2: ", sum(len(lst) for lst in fordeling.values()))
                    # Kontroller at listen regenereres korrekt, spesielt at gruppertTotal[] er oppdatert, og at priorityGroupKeys er korrekt
                    # pass på at variablene eksisterer i riktig form til å sendes til funksjonen.
            # Removing students that are fully assigned, and applications for activities that are fully assigned
            # redundant med place_student?
#            allApplications = clean_applications(allApplications, fordeling, fordelingMax, eleverMedBekreftedeAktiviteter, maxActivitiesPerStudent)
            # Students that are fully assigned, no longer have applications in allApplications (some unassigned students may have been removed)
            # Recreating groups of students with number og applications as keys
            gruppertTotal = group_student_applications("", allApplications, fordeling, maxActivitiesPerStudent, maxApplicationsPerStudent)
            # Inneholder gruppertTotal også assigned? JA!
            listOfStudentsWithFewApplications = [gruppertTotal[gruppe] for gruppe in priorityGroupKeys]
            # Checking if there still are unassigned applications for listOfStudentsWithFewApplications
                # (listOfStudentsWithFewApplications also include students with assigned activities)
            # Kontroller til slutt at denne koden ikke er redundant.
            repeatFew = False
            for i in range(len(listOfStudentsWithFewApplications)):
                for student in list(listOfStudentsWithFewApplications[i]):
                    for application in allApplications:
                        if application[0] == student:
                            repeatFew = True # kjører løkka om igjen om noen applications ikke er plassert
                    if repeatFew:
                        break
                if repeatFew:
                    break
        # end-while listOfStudentsWithFewApplications
    # end-while runWhileUnassignedRemains - ERSTATTES av bedre logikk
    
    # Itererer over keys (fordi det er en dict) og lager liste av elever som ikke er fully assigned
    # when no student needs priority assignment, this code assignes students in random order (ikke implementert?)
    # NB! elever kan ha 0-1 bekreftede aktiviteter, men ingen flere applications!
    #listOfUnassignedStudents = [student for student in eleverMedBekreftedeAktiviteter if eleverMedBekreftedeAktiviteter[student] < maxActivitiesPerStudent]
    # allApplications should now contain only applications for students that have more applications than needed to be fully assigned
    setOfUnassignedStudents = set([application[0] for application in allApplications])
    setOfStudentsWithHighPriorityApplications = set([application[0] for application in allApplications if application[3] == highPriorityString])
    #shuffle list?
    #print("Assigning students without fewApplications")
    # Bruker jeg ikke place_student her??
    for student in setOfUnassignedStudents: 
        highPriorityApplicationExists = False
        fullyAssignedGroupBreak = False
        if student in setOfStudentsWithHighPriorityApplications:
            highPriorityApplicationExists = True
        for application in list(allApplications):
            if application[0] == student:
                # Prioritizing high-priority applications
                if application[3] == highPriorityString or not highPriorityApplicationExists:
                    # Don't need to check for fully-assigned here
                    # IMPORTANT! return value needed? - Not part of a function (yet), so: NO
                    #allApplications.remove(application)
                    #place_student
                    #print("Placing student with several applications", len(allApplications))
                    allApplications, fordeling, eleverMedBekreftedeAktiviteter, allePaameldteElever = place_student(allApplications, fordeling, student, eleverMedBekreftedeAktiviteter, allePaameldteElever)
                    #fordeling[application[2].lower()].append(student)
                    # if student == "Boiko, Maksym":
                    #     print(student, "FOUND!")
                    if len(fordeling[application[2]]) == fordelingMax[application[2]]:
                        fullyAssignedGroupBreak = True # clean her? -> nei, sammen med break
                        print("Group full: ", application[2]) 
                        # group_student_applications() -> redundant, gjøres i starten av funksjonen
                        # allApplications = clean_applications() -> Kan ikke kjøre den her, inni loopene (?)
                            # eller kanskje jeg kan, siden jeg skal break ut av loopen(e)?
                        break
        # When a group is full, we need to clean applications and re-ckeckfor few allpications
        if fullyAssignedGroupBreak: 
            #allApplications = clean_applications(allApplications, fordeling, fordelingMax, eleverMedBekreftedeAktiviteter, maxActivitiesPerStudent)
            break # for student in listOfUnnassignedStudents
    #print("Number of remaining applications: ", len(allApplications))
# end while allApplications
                
#print("DEBUG EXIT_COUNTER: ", debugExitCounter)


# Creating dict with students that aren't fully assigned
studentsWithUnassignedActivities = {i:[] for i in range(1, maxActivitiesPerStudent + 1)}
#print("unassigned: ", studentsWithUnassignedActivities)
#DEBUG ==> skal skrive ut egen liste / legges til klassene
for student in eleverMedBekreftedeAktiviteter:
    for numberOfActivities in range(0, maxActivitiesPerStudent):
        if eleverMedBekreftedeAktiviteter[student] == numberOfActivities:
            if student == "Boiko, Maksym":
                print(student, "FOUND in UNASSIGNED!")
            studentsWithUnassignedActivities[maxActivitiesPerStudent - numberOfActivities].append(student)
print("studentsWithUnassignedActivities: ", len(studentsWithUnassignedActivities))

# Hver elev skal få oppfylt 2-to ønsker. 
# Elevene med 1-2 ønsker, må få sine oppfylt først.
# Det må kontrolleres om elever med 3-4 ønsker har fått sine redusert
# Deretter må resten få oppfylt sitt første ønske
# Ny kontroll må gjennomføres for hver gruppe som blir full
# Til slutt må alle gjenværende elever få oppfylt sitt andre ønske, med prioritet til "Vil mye"
# Advarsel må skrives for hver elev som ikke får oppfylt 2 ønsker

#WARNING: Systemet kan games; ved å oppgi få ønsker, er det større sjanse for å få akkurat hva du ønsker deg.
    # Elever kan risikere å ikke få en aktivitet om ønskede aktiviteter er fyllt opp.
# Hvis det blir fullt, kan elever settes opp på en tilfeldig aktivitet?



# CREATING DOCUMENTS:

#def create_word_document(file_name, headline, dictOfTuples):
def create_word_document(dictOfTuples):

    # Add data (try-except)
    for group in dictOfTuples:
        print(group) 
        file_name = group

        # Create a new Document
        doc = Document()

        # Set paper size to A4 (redundant?)
        sections = doc.sections
        for section in sections:
            section.page_width = Pt(595.3)
            section.page_height = Pt(841.9)
            section.left_margin = Pt(50)
            section.right_margin = Pt(50)
            section.top_margin = Pt(0)
            section.bottom_margin = Pt(30)

        # Add headline
        h = doc.add_heading(level=1)
        run = h.add_run(group.capitalize())
        run.font.name = 'Calibri Light'
        run.font.size = Pt(14)
        run.bold = True

        h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a table
        kolonner = 4 # I tilfelle listen er tom
#        kolonner = len(dictOfTuples[group][0]) + 1
        table = doc.add_table(rows=1, cols=kolonner) 

        # Style the table
        table.style = 'Table Grid'

        # Add the header row
        hdr_cells = table.rows[0].cells
        overskrifter = ["Navn", "Klasse", "Aktivitet"]
        for i, kolonne in enumerate(overskrifter):
            run = hdr_cells[i + 1].paragraphs[0].add_run(kolonne)
            run.font.name = 'Calibri Light'
            run.font.size = Pt(10)
            run.bold = True 

        # Adding data to table
        for x, item in enumerate(dictOfTuples[group]): # Sjekk om indeksene må justeres for header
            row_cells = table.add_row().cells
            run = row_cells[0].paragraphs[0].add_run(str(x + 1))
            for index, item2 in enumerate(item[:-1]): # Excluding the priority column
                run = row_cells[index +1].paragraphs[0].add_run(item2) 
                run.font.name = 'Calibri Light'
                run.font.size = Pt(10)
                run.bold = False


        # Set different column widths
        column_widths = [20, 150, 60, 250]
        for col_index, width in enumerate(column_widths):
            for cell in table.columns[col_index].cells:
                cell.width = Pt(width)

        # Save the document
        doc.save(f"{file_name}.docx")

# VALIDATING

# VIKTIG - Sjekk/tell kollisjon med dager
# hvilke aktivitetspar er mest utsatt?

# CREATING WORD-DOCUMENTS<<<
dictOfActivities = {activity: [] for activity in fordeling}
setOfClasses = {tup[1] for tup in finalAssembly}
dictOfClasses = {klasse: [] for klasse in setOfClasses}
for activity in fordeling:
    for student in fordeling[activity]:
        for application in finalAssembly:
            # print(application) # Funker
            if student == application[0] and activity == application[2]:  #fjerne .lower ??
                dictOfActivities[activity].append(application)
                dictOfClasses[application[1]].append(application)
                #print(application)

create_word_document(dictOfActivities)
create_word_document(dictOfClasses)
# NB! Trenger dokument over unassigned også !!